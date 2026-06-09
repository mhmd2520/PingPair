"""Group C-1 — continuous multi-segment mode tests.

Pure-Python coverage of:

1. ``SweepSegment`` / ``MultiSweepResult`` data model + helper props.
2. ``build_multi_run_report`` carry-through.
3. ``cross_segment_comparison`` table assembly.
4. Multi-segment writers — smoke tests against ``save_report`` so all
   four formats (docx, xlsx, pdf, txt) plus the v4 sidecar all land
   on disk and parse back.
5. v4 sidecar serialisation — schema_version, segments block,
   selected_case_indexes round-trip.

The Qt-side orchestration (between-segments dialog dispatch) is
deliberately not exercised here — it needs a running QApplication and
is best validated on the user's two-VM hardware.
"""

from __future__ import annotations

import json
import time
from pathlib import Path

import pytest

from pingpair.config import load_default_config
from pingpair.core.case import CaseResult
from pingpair.core.control.client import (
    MultiSweepResult,
    SweepCaseEntry,
    SweepResult,
    SweepSegment,
)
from pingpair.core.parsers.fping import FpingResult
from pingpair.core.parsers.iperf3 import IperfResult
from pingpair.core.plan import TestCase
from pingpair.core.runner import ProcSpec, RunResult
from pingpair.reporting import (
    ALL_FORMATS,
    build_multi_run_report,
    save_report,
)
from pingpair.reporting.run_report import (
    cross_segment_comparison,
    segment_summary_rows,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _rr(name: str, rc: int = 0) -> RunResult:
    return RunResult(
        spec=ProcSpec(name=name, argv=[name], cwd=Path(".")),
        returncode=rc, stdout="", stderr="",
        started_at=0.0, ended_at=1.0,
    )


def _case_result(case: TestCase, *, ok: bool = True) -> CaseResult:
    iperf = IperfResult(
        throughput_mbps=float(case.bandwidth_mbps) if ok else 0.0,
        jitter_ms=0.05 if ok else 99.0,
        packet_loss_pct=0.0 if ok else 100.0,
        raw={"end": {"sum": {}}},
    )
    fp = FpingResult(
        target="192.168.1.1",
        sent=300, received=300 if ok else 0,
        loss_pct=0.0 if ok else 100.0,
        min_ms=0.10, avg_ms=0.40, max_ms=2.10,
        elapsed_s=48.0,
    )
    return CaseResult(
        case=case,
        iperf_client=iperf,
        iperf_intervals=[],
        iperf_server_raw='{"fake": true}',
        fping=fp,
        iperf_client_run=_rr("iperf3-client"),
        fping_run=_rr("fping"),
        iperf_server_run=_rr("iperf3-server"),
        error=None if ok else "synthetic failure",
    )


def _segment(
    *, idx: int, label: str, cases: list[TestCase], all_ok: bool = True,
    status: str = "ok",
) -> SweepSegment:
    entries = [
        SweepCaseEntry(
            case=c,
            case_result=_case_result(c, ok=all_ok),
            server_iperf3_json="{}",
            server_returncode=0 if all_ok else 1,
        )
        for c in cases
    ]
    sweep = SweepResult(
        started_at=time.time() - 60,
        ended_at=time.time(),
        cases=entries,
    )
    return SweepSegment(
        segment_idx=idx, label=label, sweep=sweep,
        status=status, error="" if status == "ok" else "synthetic",
    )


def _three_segments() -> MultiSweepResult:
    """A 3-segment fixture that mirrors a small train walk-through:
    2 ok segments and 1 partial."""
    cases = [
        TestCase(index=1, payload_bytes=200, bandwidth_mbps=10, duration_s=30),
        TestCase(index=2, payload_bytes=600, bandwidth_mbps=50, duration_s=30),
    ]
    segments = [
        _segment(idx=1, label="Cab M2 ↔ M4", cases=cases),
        _segment(idx=2, label="Cab M4 ↔ M6", cases=cases),
        _segment(
            idx=3, label="Cab M6 ↔ M8",
            cases=cases, all_ok=False, status="partial",
        ),
    ]
    return MultiSweepResult(
        started_at=time.time() - 300,
        ended_at=time.time(),
        segments=segments,
        selected_case_indexes=[1, 2],
    )


@pytest.fixture
def multi() -> MultiSweepResult:
    return _three_segments()


# ---------------------------------------------------------------------------
# Data model — helpers on MultiSweepResult / SweepSegment
# ---------------------------------------------------------------------------


def test_segment_cases_ok_and_total(multi: MultiSweepResult) -> None:
    # First two segments are all-ok; third is all-failed in this fixture.
    assert multi.segments[0].cases_ok == 2
    assert multi.segments[0].cases_total == 2
    assert multi.segments[2].cases_ok == 0
    assert multi.segments[2].cases_total == 2


def test_multi_segments_ok_counts_status_ok_only(multi: MultiSweepResult) -> None:
    """``segments_ok`` counts segments where status == 'ok', regardless
    of underlying case-level errors."""
    assert multi.segments_ok == 2  # segments 1 and 2
    assert multi.segments_total == 3


def test_multi_total_cases_aggregates_across_segments(multi: MultiSweepResult) -> None:
    assert multi.total_cases == 6  # 3 segments × 2 cases
    assert multi.total_cases_ok == 4  # 2 ok segments × 2 + 0 from failed segment


def test_multi_duration_s_uses_start_end(multi: MultiSweepResult) -> None:
    assert multi.duration_s == multi.ended_at - multi.started_at
    assert multi.duration_s > 0


def test_multi_empty_segments_is_zero_duration() -> None:
    fresh = MultiSweepResult(started_at=time.time())
    assert fresh.duration_s == 0.0
    assert fresh.segments_ok == 0
    assert fresh.total_cases == 0


# ---------------------------------------------------------------------------
# build_multi_run_report
# ---------------------------------------------------------------------------


def test_build_multi_run_report_carries_metadata(multi: MultiSweepResult) -> None:
    cfg = load_default_config()
    report = build_multi_run_report(
        multi, cfg,
        run_id="UNIT_MULTI",
        metadata={"technician": "Mohamed", "customer": "TestCo"},
    )
    assert report.run_id == "UNIT_MULTI"
    assert report.segments_total == 3
    assert report.total_cases == 6
    # Metadata flows through verbatim.
    assert report.metadata == {"technician": "Mohamed", "customer": "TestCo"}
    # Selected indexes preserved.
    assert report.selected_case_indexes == [1, 2]


def test_build_multi_default_run_id_suffix_is_multisegment(
    multi: MultiSweepResult,
) -> None:
    """Without an explicit run_id, the auto-generated one ends in
    '_multisegment' so flat scans of Reports/ can tell run types apart."""
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg)
    assert report.run_id.endswith("_multisegment")


def test_build_multi_blank_label_falls_back_to_segment_n() -> None:
    cfg = load_default_config()
    cases = [
        TestCase(index=1, payload_bytes=200, bandwidth_mbps=10, duration_s=30),
    ]
    multi = MultiSweepResult(
        started_at=time.time() - 30,
        ended_at=time.time(),
        segments=[_segment(idx=7, label="", cases=cases)],
    )
    report = build_multi_run_report(multi, cfg)
    assert report.segments[0].label == "Segment 7"


# ---------------------------------------------------------------------------
# Helpers — segment_summary_rows + cross_segment_comparison
# ---------------------------------------------------------------------------


def test_segment_summary_rows_has_one_entry_per_segment(
    multi: MultiSweepResult,
) -> None:
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg)
    rows = segment_summary_rows(report)
    assert len(rows) == 3
    # Last row's status is the partial segment.
    assert rows[2][5] == "PARTIAL"


def test_cross_segment_comparison_throughput_has_segment_per_column(
    multi: MultiSweepResult,
) -> None:
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg)
    headers, rows = cross_segment_comparison(report, metric="throughput")
    # 3 case-identifying columns + 3 segment columns
    assert len(headers) == 6
    # 2 cases (the shared subset has 2 entries)
    assert len(rows) == 2
    # First row's segment-3 throughput should be 0 (synthetic failure)
    # rendered as "0.00", and the ok segments should match bandwidth.
    assert rows[0][3] == "10.00"  # segment 1, throughput == bandwidth
    assert rows[0][5] == "0.00"   # segment 3 failed → throughput 0


def test_cross_segment_comparison_unknown_metric_returns_em_dash(
    multi: MultiSweepResult,
) -> None:
    """Passing an unrecognised metric name returns dashes — defensive
    behaviour rather than raising, since the writers iterate metrics
    by string keys."""
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg)
    headers, rows = cross_segment_comparison(report, metric="banana")
    # All segment columns should be em-dash placeholders.
    for row in rows:
        assert row[3] == "—"
        assert row[4] == "—"
        assert row[5] == "—"


# ---------------------------------------------------------------------------
# Writer smoke tests — save_report dispatches on isinstance
# ---------------------------------------------------------------------------


def test_save_report_writes_all_formats_for_multi(
    multi: MultiSweepResult, tmp_path: Path,
) -> None:
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg, run_id="MULTI_ALL")
    written = save_report(
        report, dest_dir=tmp_path, basename="MULTI_ALL",
        formats=ALL_FORMATS, also_config=True, include_chart_pngs=False)
    # 4 report formats + 1 sidecar
    assert len(written) == 5
    suffixes = {p.suffix for p in written}
    assert suffixes == {".docx", ".xlsx", ".pdf", ".txt", ".json"}
    for path in written:
        assert path.exists()
        assert path.stat().st_size > 0


def test_multi_docx_contains_segment_labels(
    multi: MultiSweepResult, tmp_path: Path,
) -> None:
    """Sanity check that the docx actually mentions each segment label."""
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg, run_id="MULTI_DOCX")
    save_report(report, tmp_path, "MULTI_DOCX", ["docx"], also_config=False, include_chart_pngs=False)

    from docx import Document
    d = Document(str(tmp_path / "MULTI_DOCX" / "MULTI_DOCX.docx"))
    flat = "\n".join(p.text for p in d.paragraphs)
    flat_tables = "\n".join(
        cell.text for table in d.tables for row in table.rows for cell in row.cells
    )
    combined = flat + "\n" + flat_tables
    # All three segment labels should appear somewhere.
    assert "Cab M2 ↔ M4" in combined
    assert "Cab M4 ↔ M6" in combined
    assert "Cab M6 ↔ M8" in combined
    # The cross-segment comparison heading is included.
    assert "Cross-segment comparison" in flat


def test_multi_docx_embeds_logo_image(
    multi: MultiSweepResult, tmp_path: Path,
) -> None:
    """The multi-segment .docx carries the header logo (no appendix)."""
    from docx import Document

    from pingpair.reporting.docx_writer import write_multi_docx

    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg, run_id="MULTI_LOGO")
    dest = tmp_path / "MULTI_LOGO.docx"
    write_multi_docx(report, dest, include_appendix=False)

    doc = Document(str(dest))
    assert len(doc.inline_shapes) == 1, "expected the header logo image"


def test_multi_xlsx_has_summary_comparison_and_per_segment_sheets(
    multi: MultiSweepResult, tmp_path: Path,
) -> None:
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg, run_id="MULTI_XLSX")
    save_report(report, tmp_path, "MULTI_XLSX", ["xlsx"], also_config=False, include_chart_pngs=False)

    from openpyxl import load_workbook
    wb = load_workbook(tmp_path / "MULTI_XLSX" / "MULTI_XLSX.xlsx")
    sheet_names = wb.sheetnames
    # Required fixed sheets.
    assert "Summary" in sheet_names
    assert "Run info" in sheet_names
    assert "Throughput (Mbps)" in sheet_names
    assert "Avg Latency (ms)" in sheet_names
    assert "Packet Loss (%)" in sheet_names
    # One detail sheet per segment (auto-generated names, just check count).
    detail_sheets = [n for n in sheet_names if n.startswith("Seg ")]
    assert len(detail_sheets) == 3


def test_multi_xlsx_sheet_name_sanitised_for_long_label(tmp_path: Path) -> None:
    """Labels with forbidden chars or >31 chars don't crash openpyxl."""
    cfg = load_default_config()
    long_label = "Cab/M2-A:Section[7] very long label that exceeds limit"
    cases = [TestCase(index=1, payload_bytes=200, bandwidth_mbps=10, duration_s=30)]
    multi = MultiSweepResult(
        started_at=time.time() - 30,
        ended_at=time.time(),
        segments=[_segment(idx=1, label=long_label, cases=cases)],
    )
    report = build_multi_run_report(multi, cfg, run_id="LONG_SHEET")
    # Must not raise.
    save_report(report, tmp_path, "LONG_SHEET", ["xlsx"], also_config=False, include_chart_pngs=False)
    out = tmp_path / "LONG_SHEET" / "LONG_SHEET.xlsx"
    assert out.exists()
    from openpyxl import load_workbook
    wb = load_workbook(out)
    # Find the sanitised sheet — must be ≤ 31 chars and no forbidden chars.
    detail_sheets = [n for n in wb.sheetnames if n.startswith("Seg ")]
    assert len(detail_sheets) == 1
    name = detail_sheets[0]
    assert len(name) <= 31
    for ch in "[]:*?/\\":
        assert ch not in name


def test_multi_txt_contains_all_metric_comparison_tables(
    multi: MultiSweepResult, tmp_path: Path,
) -> None:
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg, run_id="MULTI_TXT")
    save_report(report, tmp_path, "MULTI_TXT", ["txt"], also_config=False, include_chart_pngs=False)
    text = (tmp_path / "MULTI_TXT" / "MULTI_TXT.txt").read_text("utf-8")
    assert "Cross-segment comparison: Throughput Received (Mbps)" in text
    assert "Cross-segment comparison: Average Latency (ms)" in text
    assert "Cross-segment comparison: Packet Loss (%)" in text


# ---------------------------------------------------------------------------
# Schema v4 sidecar
# ---------------------------------------------------------------------------


def test_multi_config_sidecar_schema_version_is_5(
    multi: MultiSweepResult, tmp_path: Path,
) -> None:
    """Group F (Q1, 2026-05-16) bumped multi-segment sidecar from v4 to v5
    with the additive ``gateway`` + ``nic_override`` keys. Q2 also
    renamed the file extension from ``.config.json`` to ``.json``."""
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg, run_id="V5")
    save_report(report, tmp_path, "V5", formats=[], also_config=True, include_chart_pngs=False)
    sidecar = tmp_path / "V5" / "V5.json"
    data = json.loads(sidecar.read_text(encoding="utf-8"))
    assert data["schema_version"] == 5
    assert "segments" in data
    assert len(data["segments"]) == 3
    # Group F additive keys must be present even on a default profile.
    assert "gateway" in data
    assert "nic_override" in data
    # Each segment in the sidecar carries label + status + cases.
    seg0 = data["segments"][0]
    assert seg0["label"] == "Cab M2 ↔ M4"
    assert seg0["status"] == "ok"
    assert seg0["cases_total"] == 2
    assert isinstance(seg0["cases"], list)


def test_multi_config_sidecar_carries_selected_case_indexes(
    multi: MultiSweepResult, tmp_path: Path,
) -> None:
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg, run_id="V4_SEL")
    save_report(report, tmp_path, "V4_SEL", formats=[], also_config=True, include_chart_pngs=False)
    data = json.loads(
        (tmp_path / "V4_SEL" / "V4_SEL.json").read_text("utf-8")
    )
    assert data["selected_case_indexes"] == [1, 2]


def test_multi_config_sidecar_top_level_aggregates_match_dataclass(
    multi: MultiSweepResult, tmp_path: Path,
) -> None:
    cfg = load_default_config()
    report = build_multi_run_report(multi, cfg, run_id="V4_AGG")
    save_report(
        report, tmp_path, "V4_AGG",
        formats=[], also_config=True, include_chart_pngs=False,
    )
    data = json.loads((tmp_path / "V4_AGG" / "V4_AGG.json").read_text("utf-8"))
    # Top-level aggregates must match the dataclass properties exactly.
    assert data["segments_total"] == report.segments_total
    assert data["segments_ok"] == report.segments_ok
    assert data["total_cases"] == report.total_cases
    assert data["total_cases_ok"] == report.total_cases_ok
    # And the per-segment block carries one entry per segment.
    assert len(data["segments"]) == report.segments_total
