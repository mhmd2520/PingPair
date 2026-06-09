"""Smoke tests for the report writers.

We synthesise a minimal SweepResult + AppConfig, run each writer to a
temp directory, and check that:

* the file is created and non-empty,
* it parses back / has the right magic bytes,
* the test case's metric values appear in the rendered output.
"""

from __future__ import annotations

import json
import time
import zipfile
from pathlib import Path

import pytest

from pingpair.config import load_default_config
from pingpair.core.case import CaseResult
from pingpair.core.control.client import SweepCaseEntry, SweepResult
from pingpair.core.parsers.fping import FpingResult
from pingpair.core.parsers.iperf3 import IperfResult
from pingpair.core.plan import TestCase
from pingpair.core.runner import RunResult, fping_spec, iperf3_client_spec
from pingpair.reporting import (
    ALL_FORMATS,
    build_run_report,
    save_report,
    unique_basename,
)


# ---------------------------------------------------------------------------
# Fixtures
# ---------------------------------------------------------------------------


def _fake_run_result(name: str, rc: int = 0) -> RunResult:
    """Build a RunResult with placeholder timing."""
    from pingpair.core.runner import ProcSpec
    return RunResult(
        spec=ProcSpec(name=name, argv=[name], cwd=Path(".")),
        returncode=rc,
        stdout="",
        stderr="",
        started_at=0.0,
        ended_at=1.0,
    )


def _fake_case_result(case: TestCase) -> CaseResult:
    """Build a CaseResult that looks like a clean run on a virtual switch."""
    iperf = IperfResult(
        throughput_mbps=float(case.bandwidth_mbps),  # near-perfect
        jitter_ms=0.05,
        packet_loss_pct=0.0,
        raw={"end": {"sum": {}}},
    )
    fp = FpingResult(
        target="192.168.1.1",
        sent=3000,
        received=3000,
        loss_pct=0.0,
        min_ms=0.10,
        avg_ms=0.40,
        max_ms=2.10,
        elapsed_s=48.0,
    )
    return CaseResult(
        case=case,
        iperf_client=iperf,
        iperf_intervals=[],
        iperf_server_raw='{"fake": true}',
        fping=fp,
        iperf_client_run=_fake_run_result("iperf3-client"),
        fping_run=_fake_run_result("fping"),
        iperf_server_run=_fake_run_result("iperf3-server"),
        error=None,
    )


def _fake_sweep_result() -> SweepResult:
    """Build a 3-case fake sweep so tests are quick but still cover real shapes."""
    cases = [
        TestCase(index=1, payload_bytes=200, bandwidth_mbps=10, duration_s=30),
        TestCase(index=2, payload_bytes=600, bandwidth_mbps=50, duration_s=30),
        TestCase(index=3, payload_bytes=1300, bandwidth_mbps=90, duration_s=30),
    ]
    entries = [
        SweepCaseEntry(
            case=c,
            case_result=_fake_case_result(c),
            server_iperf3_json='{"end": {}}',
            server_returncode=0,
        )
        for c in cases
    ]
    return SweepResult(
        started_at=time.time() - 100,
        ended_at=time.time(),
        cases=entries,
    )


@pytest.fixture
def sweep() -> SweepResult:
    return _fake_sweep_result()


# ---------------------------------------------------------------------------
# Builder
# ---------------------------------------------------------------------------


def test_build_run_report_carries_per_case_metrics(sweep: SweepResult) -> None:
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UNIT_TEST_RUN")
    assert report.run_id == "UNIT_TEST_RUN"
    assert report.cases_total == 3
    assert report.cases_ok == 3
    assert report.cases[0].payload_bytes == 200
    assert report.cases[0].bandwidth_mbps_pushed == 10
    assert report.cases[0].throughput_mbps_received == 10.0
    assert report.cases[2].max_latency_ms == 2.10


# ---------------------------------------------------------------------------
# Writers
# ---------------------------------------------------------------------------


def test_save_report_writes_all_requested_formats(
    sweep: SweepResult, tmp_path: Path
) -> None:
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="ALL_FMTS")
    written = save_report(
        report,
        dest_dir=tmp_path,
        basename="ALL_FMTS",
        formats=ALL_FORMATS,
        also_config=True, include_chart_pngs=False)
    # 4 formats + 1 config sidecar
    assert len(written) == 5
    expected_suffixes = {".docx", ".xlsx", ".pdf", ".txt", ".json"}
    assert {p.suffix for p in written} == expected_suffixes
    for path in written:
        assert path.exists()
        assert path.stat().st_size > 0


def test_xlsx_summary_sheet_carries_throughput(
    sweep: SweepResult, tmp_path: Path
) -> None:
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_XLSX")
    save_report(report, tmp_path, "UT_XLSX", ["xlsx"], also_config=False, include_chart_pngs=False)

    from openpyxl import load_workbook
    wb = load_workbook(tmp_path / "UT_XLSX" / "UT_XLSX.xlsx")
    assert "Summary" in wb.sheetnames
    assert "Detail" in wb.sheetnames
    assert "Run info" in wb.sheetnames
    sm = wb["Summary"]
    # Header row
    assert sm.cell(1, 1).value == "#"
    assert sm.cell(1, 4).value == "Throughput Received (Mbps)"
    # First data row matches the synthetic case (case_idx=1, payload=200, bw=10, throughput=10.0)
    assert sm.cell(2, 1).value == 1
    assert sm.cell(2, 2).value == 200
    assert sm.cell(2, 3).value == 10
    assert sm.cell(2, 4).value == 10.0


def test_docx_round_trips(sweep: SweepResult, tmp_path: Path) -> None:
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_DOCX")
    save_report(report, tmp_path, "UT_DOCX", ["docx"], also_config=False, include_chart_pngs=False)

    # python-docx round-trip: re-open and confirm the title is there.
    from docx import Document
    d = Document(str(tmp_path / "UT_DOCX" / "UT_DOCX.docx"))
    text_blob = "\n".join(p.text for p in d.paragraphs)
    assert "PingPair" in text_blob
    assert "Performance Metrics" in text_blob
    # Verify a known case value appears in *some* table cell.
    flat_table_text = "\n".join(
        cell.text for table in d.tables for row in table.rows for cell in row.cells
    )
    assert "200" in flat_table_text
    assert "10" in flat_table_text


def test_txt_contains_summary_table(sweep: SweepResult, tmp_path: Path) -> None:
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_TXT")
    save_report(report, tmp_path, "UT_TXT", ["txt"], also_config=False, include_chart_pngs=False)

    text = (tmp_path / "UT_TXT" / "UT_TXT.txt").read_text(encoding="utf-8")
    assert "Performance Metrics" in text
    assert "Per-case detail" in text
    # Throughput-like value
    assert "10.00" in text or "10.0" in text


def test_pdf_starts_with_pdf_magic(sweep: SweepResult, tmp_path: Path) -> None:
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_PDF")
    save_report(report, tmp_path, "UT_PDF", ["pdf"], also_config=False, include_chart_pngs=False)

    head = (tmp_path / "UT_PDF" / "UT_PDF.pdf").read_bytes()[:5]
    assert head == b"%PDF-"


def test_config_sidecar_round_trips(
    sweep: SweepResult, tmp_path: Path
) -> None:
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_CFG")
    save_report(report, tmp_path, "UT_CFG", [], also_config=True, include_chart_pngs=False)

    data = json.loads((tmp_path / "UT_CFG" / "UT_CFG.json").read_text(encoding="utf-8"))
    assert data["run_id"] == "UT_CFG"
    assert data["cases_total"] == 3
    assert data["cases_ok"] == 3
    assert len(data["cases"]) == 3
    assert data["cases"][0]["payload_bytes"] == 200
    assert data["cases"][0]["throughput_mbps_received"] == 10.0


def test_save_report_creates_per_sweep_subfolder(
    sweep: SweepResult, tmp_path: Path
) -> None:
    """Each sweep should land in its own Reports/<basename>/ subfolder."""
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_FOLDER")
    written = save_report(report, tmp_path, "UT_FOLDER", ["docx"], also_config=True, include_chart_pngs=False)

    sweep_dir = tmp_path / "UT_FOLDER"
    assert sweep_dir.is_dir(), "save_report should create a per-sweep subfolder"
    # Both files should be inside that subfolder.
    for path in written:
        assert path.parent == sweep_dir
    assert (sweep_dir / "UT_FOLDER.docx").exists()
    assert (sweep_dir / "UT_FOLDER.json").exists()


def test_fmt_duration_formats_minutes_and_seconds() -> None:
    from pingpair.reporting.run_report import fmt_duration
    assert fmt_duration(0) == "0s"
    assert fmt_duration(30) == "30s"
    assert fmt_duration(60) == "1m 0s"
    assert fmt_duration(950.6) == "15m 51s"
    assert fmt_duration(965.2) == "16m 5s"


# ---------------------------------------------------------------------------
# Cable length (Run-tab optional field) — flows into every report
# ---------------------------------------------------------------------------


def test_cable_length_in_metadata_rows(sweep: SweepResult) -> None:
    from pingpair.reporting.run_report import metadata_rows
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_CL", cable_length_m="12.50")
    assert report.cable_length_m == "12.50"
    assert ("Cable length (m)", "12.50") in metadata_rows(report)


def test_cable_length_blank_omitted_from_metadata_rows(sweep: SweepResult) -> None:
    from pingpair.reporting.run_report import metadata_rows
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_CL0")
    labels = [label for label, _ in metadata_rows(report)]
    assert "Cable length (m)" not in labels


def test_cable_length_in_sidecar(sweep: SweepResult, tmp_path: Path) -> None:
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_CLJSON", cable_length_m="3.14")
    save_report(report, tmp_path, "UT_CLJSON", [], also_config=True, include_chart_pngs=False)
    data = json.loads(
        (tmp_path / "UT_CLJSON" / "UT_CLJSON.json").read_text(encoding="utf-8")
    )
    assert data["cable_length_m"] == "3.14"


# ---------------------------------------------------------------------------
# xlsx Analysis charts carry explicit markers (#2) — so a SINGLE-case run
# (one data point) is still visible instead of an empty-looking chart.
# ---------------------------------------------------------------------------


def _single_case_sweep(sweep: SweepResult) -> SweepResult:
    return SweepResult(
        cases=sweep.cases[:1],
        started_at=sweep.started_at,
        ended_at=sweep.ended_at,
    )


def test_xlsx_charts_have_markers_single_case(sweep: SweepResult) -> None:
    from openpyxl import Workbook

    from pingpair.reporting.appendix import append_to_xlsx

    cfg = load_default_config()
    report = build_run_report(_single_case_sweep(sweep), cfg, run_id="UT_MK1")
    wb = Workbook()
    append_to_xlsx(wb, report, is_multi=False)
    charts = wb["Analysis"]._charts
    assert charts, "expected native Excel charts in the Analysis sheet"
    for chart in charts:
        for series in chart.series:
            assert series.marker is not None
            assert series.marker.symbol == "circle"


# ---------------------------------------------------------------------------
# unique_basename — Group C-1 follow-up 2026-05-11
# ---------------------------------------------------------------------------


def test_unique_basename_returns_original_when_folder_absent(tmp_path: Path) -> None:
    """First save with a brand-new pattern keeps the typed name."""
    assert unique_basename(tmp_path, "test") == "test"


def test_unique_basename_appends_2_when_folder_exists(tmp_path: Path) -> None:
    (tmp_path / "test").mkdir()
    assert unique_basename(tmp_path, "test") == "test_2"


def test_unique_basename_skips_past_existing_suffixes(tmp_path: Path) -> None:
    """Both base + _2 exist → return _3 (not _2 reused)."""
    (tmp_path / "M2-M4-baseline").mkdir()
    (tmp_path / "M2-M4-baseline_2").mkdir()
    assert unique_basename(tmp_path, "M2-M4-baseline") == "M2-M4-baseline_3"


def test_unique_basename_returns_base_when_only_suffix_present(tmp_path: Path) -> None:
    """If the user manually deleted the original 'test' folder but left
    'test_2' around, the un-suffixed slot is reclaimed rather than
    forcing the user further down the suffix chain."""
    (tmp_path / "test_2").mkdir()
    # 'test' itself doesn't exist → take it.
    assert unique_basename(tmp_path, "test") == "test"


def test_unique_basename_also_avoids_collision_with_files(tmp_path: Path) -> None:
    """A regular file with the target name still counts as a collision."""
    (tmp_path / "test").write_text("not a folder", encoding="utf-8")
    assert unique_basename(tmp_path, "test") == "test_2"


def test_unique_basename_handles_default_multisegment_suffix(tmp_path: Path) -> None:
    """The multi-segment save helper appends '_multisegment' before
    calling unique_basename — verify the chain works."""
    (tmp_path / "PingPair_2026-05-11_125954_multisegment").mkdir()
    out = unique_basename(tmp_path, "PingPair_2026-05-11_125954_multisegment")
    assert out == "PingPair_2026-05-11_125954_multisegment_2"


def test_unknown_format_raises(sweep: SweepResult, tmp_path: Path) -> None:
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_BAD")
    with pytest.raises(ValueError, match="unknown report format"):
        save_report(report, tmp_path, "UT_BAD", ["mp4"], also_config=False, include_chart_pngs=False)  # type: ignore[list-item]


# ---------------------------------------------------------------------------
# Report-header logo (2026-06-09) — the PingPair mark is embedded at the top
# of every Word / PDF report, rendered headless from branding.ICON_SVG.
# ---------------------------------------------------------------------------


def test_logo_png_bytes_are_a_valid_png() -> None:
    """branding.logo_png_bytes renders a real PNG without a QApplication."""
    from pingpair.branding import logo_png_bytes

    png = logo_png_bytes(64)
    assert png.startswith(b"\x89PNG\r\n\x1a\n"), "expected PNG magic bytes"
    assert len(png) > 100


def test_docx_embeds_logo_image(sweep: SweepResult, tmp_path: Path) -> None:
    """The single-sweep .docx carries exactly one image — the header logo —
    when the chart appendix is disabled (the appendix adds its own images)."""
    from docx import Document

    from pingpair.reporting.docx_writer import write_docx

    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_LOGO")
    dest = tmp_path / "UT_LOGO.docx"
    write_docx(report, dest, include_appendix=False)

    doc = Document(str(dest))
    assert len(doc.inline_shapes) == 1, "expected the header logo image"


def test_pdf_embeds_logo_image(sweep: SweepResult, tmp_path: Path) -> None:
    """The single-sweep .pdf embeds the header logo as an image XObject
    (appendix disabled, so the logo is the only embedded image)."""
    from pingpair.reporting.pdf_writer import write_pdf

    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_PDF_LOGO")
    dest = tmp_path / "UT_PDF_LOGO.pdf"
    write_pdf(report, dest, include_appendix=False)

    data = dest.read_bytes()
    assert data[:5] == b"%PDF-"
    assert b"/XObject" in data, "expected the header logo image XObject in the PDF"


def test_logo_render_failure_is_a_clean_noop(monkeypatch) -> None:
    """The failure-tolerance contract: when rendering yields no bytes, the logo
    is omitted cleanly — the PDF flowable is None and the docx gains no stray
    paragraph or image (so a render failure can never break a report)."""
    from docx import Document

    from pingpair.reporting import _logo

    monkeypatch.setattr(_logo, "_logo_png", lambda *a, **k: None)

    assert _logo.pdf_logo_flowable() is None

    doc = Document()
    _logo.add_docx_logo(doc)
    assert len(doc.inline_shapes) == 0, "no image should be embedded on render failure"
    assert len(doc.paragraphs) == 0, "no stray empty paragraph should be left behind"


def test_metadata_flows_through_every_format(
    sweep: SweepResult, tmp_path: Path
) -> None:
    """Test-record metadata should land in docx/xlsx/pdf/txt/<basename>.json."""
    cfg = load_default_config()
    metadata = {
        "technician": "Mohamed Khaled",
        "customer": "TestCo Train Site",
        "hardware_sn": "MCG-CAB-M2-9876",
        "environment": "Lab bench, 2-VM staging",
        "record_id": "TR-2026-0042",
    }
    report = build_run_report(sweep, cfg, run_id="UT_META", metadata=metadata)
    save_report(report, tmp_path, "UT_META", ["docx", "xlsx", "txt"], also_config=True, include_chart_pngs=False)

    # docx — the title-block table cells should contain every value.
    from docx import Document
    doc = Document(str(tmp_path / "UT_META" / "UT_META.docx"))
    docx_text = "\n".join(
        cell.text for table in doc.tables for row in table.rows for cell in row.cells
    )
    for value in metadata.values():
        assert value in docx_text, f"{value!r} missing from docx"

    # xlsx — Run-info sheet should have rows for each populated field.
    from openpyxl import load_workbook
    wb = load_workbook(tmp_path / "UT_META" / "UT_META.xlsx")
    run_info_text = " ".join(
        str(c.value or "")
        for row in wb["Run info"].iter_rows()
        for c in row
    )
    for value in metadata.values():
        assert value in run_info_text, f"{value!r} missing from xlsx Run info"

    # txt — header should list every metadata line.
    txt = (tmp_path / "UT_META" / "UT_META.txt").read_text(encoding="utf-8")
    for value in metadata.values():
        assert value in txt, f"{value!r} missing from txt"

    # Sidecar (.json) — top-level "metadata" key should preserve the dict.
    cfg_data = json.loads((tmp_path / "UT_META" / "UT_META.json").read_text(encoding="utf-8"))
    assert cfg_data["metadata"] == metadata
    # schema_version history: 2 → 3 added ``selected_case_indexes`` (Group B);
    # 3 → 5 added ``gateway`` + ``nic_override`` (Group F / Q1, 2026-05-16).
    # If you bump it again, update reporting/config_writer.py::_serialise()
    # and this assertion together so they don't drift.
    assert cfg_data["schema_version"] == 5


def test_empty_metadata_is_omitted_from_visible_outputs(
    sweep: SweepResult, tmp_path: Path
) -> None:
    """Blank fields should NOT show up as 'Field: ' empty rows."""
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_EMPTY", metadata={})
    save_report(report, tmp_path, "UT_EMPTY", ["txt"], also_config=False, include_chart_pngs=False)
    txt = (tmp_path / "UT_EMPTY" / "UT_EMPTY.txt").read_text(encoding="utf-8")
    for label in ("Technician:", "Customer / Site:", "Hardware S/N:"):
        assert label not in txt, f"{label!r} unexpectedly appears in empty-metadata report"


# ---------------------------------------------------------------------------
# Group F (Q1 task #13, 2026-05-16): sidecar schema v5 — additive gateway
# + nic_override snapshot. v3/v4 readers parse cleanly.
# ---------------------------------------------------------------------------


def test_sidecar_v5_records_gateway_none_for_canonical_setup(
    sweep: SweepResult, tmp_path: Path,
) -> None:
    """defaults.json ships gateway: null -> sidecar emits gateway: null."""
    import json
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_V5_NULL")
    save_report(report, tmp_path, "UT_V5_NULL", ["docx"])
    sidecar = tmp_path / "UT_V5_NULL" / "UT_V5_NULL.json"
    data = json.loads(sidecar.read_text())
    assert data["schema_version"] == 5
    assert "gateway" in data
    assert data["gateway"] is None
    assert "nic_override" in data
    assert data["nic_override"] is None


def test_sidecar_v5_records_gateway_when_profile_has_one(
    sweep: SweepResult, tmp_path: Path,
) -> None:
    """A profile with a gateway set surfaces in the sidecar verbatim."""
    import json
    cfg = load_default_config()
    cfg.network.gateway = "192.168.1.254"  # type: ignore[assignment]
    report = build_run_report(sweep, cfg, run_id="UT_V5_GW")
    save_report(report, tmp_path, "UT_V5_GW", ["docx"])
    sidecar = tmp_path / "UT_V5_GW" / "UT_V5_GW.json"
    data = json.loads(sidecar.read_text())
    assert data["schema_version"] == 5
    assert data["gateway"] == "192.168.1.254"


def test_sidecar_v5_loads_via_existing_v3_v4_loader(
    sweep: SweepResult, tmp_path: Path,
) -> None:
    """The Analysis tab's sidecar_loader (written for v1-v4) must parse
    v5 cleanly because the new fields are purely additive."""
    from pingpair.analysis import load_sidecar
    cfg = load_default_config()
    report = build_run_report(sweep, cfg, run_id="UT_V5_BACKCOMPAT")
    save_report(report, tmp_path, "UT_V5_BACKCOMPAT", ["docx"])
    sidecar = tmp_path / "UT_V5_BACKCOMPAT" / "UT_V5_BACKCOMPAT.json"
    run = load_sidecar(sidecar)
    assert run.schema_version == 5
    assert run.run_id == "UT_V5_BACKCOMPAT"
    # The cases still flatten into a Series like a v3 single-segment run.
    assert len(run.series) == 1
