"""Render a :class:`RunReport` as a Word document.

Layout (matches the project's Table-1.PNG and the Test Procedure record):

* Title page header with run ID + timestamps + IPs + software versions.
* "Performance Metrics" section with the 8-column results table.
* "Per-case detail" section: one row per case with full forensics
  (status, error, return codes, exact CLI strings).
* Footer note with PingPair version.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt

from ._logo import add_docx_logo
from .run_report import (
    CaseMetrics,
    MultiSegmentRunReport,
    RunReport,
    SEGMENT_SUMMARY_HEADERS,
    cross_segment_comparison,
    fmt_duration,
    metadata_rows,
    segment_summary_rows,
)


_HEADERS_MAIN: tuple[str, ...] = (
    "Payload Size\n(Bytes)",
    "Bandwidth\n(MBits Per Second)\n– Pushed Traffic",
    "Throughput\n(MBits Per Second)\n– Received Traffic",
    "Jitter\n(msec)",
    "Packet Loss\n(%)",
    "Minimum Latency\n(msec)",
    "Average Latency\n(msec)",
    "Maximum Latency\n(msec)",
)


def _fmt(v: float | None, spec: str) -> str:
    return "—" if v is None else format(v, spec)


def write_docx(
    report: RunReport,
    dest: Path,
    *,
    include_appendix: bool = False,
) -> None:
    """Write the .docx report to ``dest``.

    When ``include_appendix`` is True, the file gets an Analysis appendix
    (per-metric summary stats) appended just before save (#13).
    """
    doc = Document()

    # ---- Title block ------------------------------------------------------
    add_docx_logo(doc)
    title = doc.add_heading("PingPair — LAN Characterization Report", level=0)
    title.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

    info = doc.add_table(rows=0, cols=2)
    info.style = "Light Grid Accent 1"
    base_rows = [
        ("Run ID", report.run_id),
        ("Started", report.started_at.strftime("%Y-%m-%d %H:%M:%S")),
        ("Finished", report.ended_at.strftime("%Y-%m-%d %H:%M:%S")),
        ("Duration", fmt_duration(report.duration_s)),
        ("Server IP", report.server_ip),
        ("Client IP", report.client_ip),
        ("Protocol", report.protocol.upper()),
        ("Cases ok", f"{report.cases_ok} / {report.cases_total}"),
        ("fping version", report.fping_version),
        ("iperf3 version", report.iperf3_version),
        ("PingPair version", report.app_version),
    ]
    # Test-record metadata is added if any fields are populated.
    base_rows.extend(metadata_rows(report))
    for label, value in base_rows:
        row = info.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_paragraph("")

    # ---- Performance metrics table (Table-1.PNG shape) -------------------
    doc.add_heading("Performance Metrics", level=1)
    doc.add_paragraph(
        "Steady-state network with no external simulated or video traffic. "
        "QoS: COS set to 0."
    )

    table = doc.add_table(rows=1, cols=len(_HEADERS_MAIN))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(_HEADERS_MAIN):
        hdr[i].text = h
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for c in report.cases:
        row = table.add_row().cells
        values = (
            str(c.payload_bytes),
            str(c.bandwidth_mbps_pushed),
            _fmt(c.throughput_mbps_received, ".2f"),
            _fmt(c.jitter_ms, ".3f"),
            _fmt(c.packet_loss_pct, ".3f"),
            _fmt(c.min_latency_ms, ".2f"),
            _fmt(c.avg_latency_ms, ".2f"),
            _fmt(c.max_latency_ms, ".2f"),
        )
        for i, v in enumerate(values):
            row[i].text = v
            for paragraph in row[i].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # ---- Per-case detail --------------------------------------------------
    doc.add_paragraph("")
    doc.add_heading("Per-case detail", level=1)
    doc.add_paragraph(
        "Forensics for each case: status, error notes if any, subprocess "
        "return codes, and the exact CLI strings used. "
        "Useful for re-running a single case from a terminal."
    )

    for c in report.cases:
        doc.add_heading(
            f"Case #{c.case_idx:02d} — payload {c.payload_bytes} B, "
            f"bandwidth {c.bandwidth_mbps_pushed} Mbps "
            f"({c.duration_s} s, {report.protocol.upper()})",
            level=2,
        )
        detail = doc.add_table(rows=0, cols=2)
        detail.style = "Light Grid Accent 1"
        for label, value in (
            ("Status", c.status),
            ("Error", c.error or "—"),
            ("iperf3 client rc", str(c.iperf3_client_rc) if c.iperf3_client_rc is not None else "—"),
            ("iperf3 server rc", str(c.iperf3_server_rc) if c.iperf3_server_rc is not None else "—"),
            ("fping rc", str(c.fping_rc) if c.fping_rc is not None else "—"),
            ("Throughput received", _fmt(c.throughput_mbps_received, ".3f") + " Mbps"),
            ("Jitter", _fmt(c.jitter_ms, ".3f") + " ms"),
            ("Loss", _fmt(c.packet_loss_pct, ".3f") + " %"),
            ("Min / Avg / Max latency",
             f"{_fmt(c.min_latency_ms, '.2f')} / "
             f"{_fmt(c.avg_latency_ms, '.2f')} / "
             f"{_fmt(c.max_latency_ms, '.2f')} ms"),
            ("iperf3 client cmd", c.iperf3_client_cmd),
            ("iperf3 server cmd", c.iperf3_server_cmd),
            ("fping cmd", c.fping_cmd),
        ):
            row = detail.add_row().cells
            row[0].text = label
            row[1].text = value

    # ---- Footer ----------------------------------------------------------
    doc.add_paragraph("")
    footer = doc.add_paragraph()
    footer.alignment = 1  # CENTER
    run = footer.add_run(f"Generated by PingPair {report.app_version}")
    run.italic = True
    run.font.size = Pt(8)

    dest.parent.mkdir(parents=True, exist_ok=True)
    if include_appendix:
        import tempfile
        from pathlib import Path as _Path
        from .appendix import append_to_docx
        with tempfile.TemporaryDirectory(prefix="pingtool_appendix_") as tmp:
            append_to_docx(doc, report, is_multi=False, tmp_dir=_Path(tmp))
            doc.save(str(dest))
    else:
        doc.save(str(dest))


# ---------------------------------------------------------------------------
# Group C-1: multi-segment writer
# ---------------------------------------------------------------------------


def _case_row_values(c: CaseMetrics) -> tuple[str, ...]:
    """The 8-column row used in both single and multi-segment Performance Metrics tables."""
    return (
        str(c.payload_bytes),
        str(c.bandwidth_mbps_pushed),
        _fmt(c.throughput_mbps_received, ".2f"),
        _fmt(c.jitter_ms, ".3f"),
        _fmt(c.packet_loss_pct, ".3f"),
        _fmt(c.min_latency_ms, ".2f"),
        _fmt(c.avg_latency_ms, ".2f"),
        _fmt(c.max_latency_ms, ".2f"),
    )


def write_multi_docx(
    report: MultiSegmentRunReport,
    dest: Path,
    *,
    include_appendix: bool = False,
) -> None:
    """Write a multi-segment .docx report to ``dest``.

    Layout: title → run-level info table → segments summary →
    per-segment Performance Metrics tables → three cross-segment
    comparison tables (Throughput / Avg Latency / Loss) → footer.
    The per-case detail forensic dump is omitted on purpose — for
    multi-segment runs the detail would balloon to N × 20 sub-tables
    and obscure the headline comparison view. Forensics are still in
    the .json sidecar.
    """
    doc = Document()

    # ---- Title block ------------------------------------------------------
    add_docx_logo(doc)
    title = doc.add_heading(
        "PingPair — LAN Characterization Report (Multi-Segment)", level=0
    )
    title.alignment = 1  # CENTER

    info = doc.add_table(rows=0, cols=2)
    info.style = "Light Grid Accent 1"
    base_rows = [
        ("Run ID", report.run_id),
        ("Started", report.started_at.strftime("%Y-%m-%d %H:%M:%S")),
        ("Finished", report.ended_at.strftime("%Y-%m-%d %H:%M:%S")),
        ("Total duration", fmt_duration(report.duration_s)),
        ("Server IP", report.server_ip),
        ("Client IP", report.client_ip),
        ("Protocol", report.protocol.upper()),
        ("Segments ok", f"{report.segments_ok} / {report.segments_total}"),
        ("Total cases ok",
         f"{report.total_cases_ok} / {report.total_cases}"),
        ("fping version", report.fping_version),
        ("iperf3 version", report.iperf3_version),
        ("PingPair version", report.app_version),
    ]
    base_rows.extend(metadata_rows(report))
    for label, value in base_rows:
        row = info.add_row().cells
        row[0].text = label
        row[1].text = str(value)

    doc.add_paragraph("")

    # ---- Segments summary -----------------------------------------------
    doc.add_heading("Segments summary", level=1)
    doc.add_paragraph(
        "One row per segment in the order it ran. Status column is OK "
        "(all cases passed), PARTIAL (some cases errored but the segment "
        "completed), or FAILED (segment couldn't complete, e.g. TCP drop)."
    )
    table = doc.add_table(rows=1, cols=len(SEGMENT_SUMMARY_HEADERS))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(SEGMENT_SUMMARY_HEADERS):
        hdr[i].text = h
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    for row_values in segment_summary_rows(report):
        row = table.add_row().cells
        for i, v in enumerate(row_values):
            row[i].text = v

    # ---- Per-segment Performance Metrics tables -------------------------
    doc.add_paragraph("")
    doc.add_heading("Per-segment performance metrics", level=1)
    doc.add_paragraph(
        "The 8-column metric grid for each segment. Identical shape to "
        "the single-sweep report — handy if you want to compare a "
        "specific segment against a historic baseline."
    )

    for seg in report.segments:
        doc.add_heading(
            f"Segment #{seg.segment_idx} — {seg.label} "
            f"({fmt_duration(seg.duration_s)}, {seg.status.upper()})",
            level=2,
        )
        if seg.error:
            err_para = doc.add_paragraph()
            err_run = err_para.add_run(f"Note: {seg.error}")
            err_run.italic = True
            err_run.font.size = Pt(9)

        seg_table = doc.add_table(rows=1, cols=len(_HEADERS_MAIN))
        seg_table.style = "Light Grid Accent 1"
        hdr = seg_table.rows[0].cells
        for i, h in enumerate(_HEADERS_MAIN):
            hdr[i].text = h
            for paragraph in hdr[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
            hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        for c in seg.cases:
            row = seg_table.add_row().cells
            for i, v in enumerate(_case_row_values(c)):
                row[i].text = v
                for paragraph in row[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(10)

    # ---- Cross-segment comparison ---------------------------------------
    doc.add_paragraph("")
    doc.add_heading("Cross-segment comparison", level=1)
    doc.add_paragraph(
        "Each table below shows one metric across every segment, "
        "letting you spot regressions or sudden drops between car-pairs "
        "at a glance. Cells read '—' when the case didn't run in that "
        "segment (e.g. a partial / failed segment cut short)."
    )

    for metric_id, metric_label in (
        ("throughput", "Throughput Received (Mbps)"),
        ("avg_latency", "Average Latency (ms)"),
        ("loss", "Packet Loss (%)"),
    ):
        doc.add_heading(metric_label, level=2)
        headers, rows = cross_segment_comparison(report, metric=metric_id)
        cmp_table = doc.add_table(rows=1, cols=len(headers))
        cmp_table.style = "Light Grid Accent 1"
        hdr_cells = cmp_table.rows[0].cells
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        for row_values in rows:
            row = cmp_table.add_row().cells
            for i, v in enumerate(row_values):
                row[i].text = v
                for paragraph in row[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)
        doc.add_paragraph("")

    # ---- Footer ---------------------------------------------------------
    footer = doc.add_paragraph()
    footer.alignment = 1
    run = footer.add_run(f"Generated by PingPair {report.app_version}")
    run.italic = True
    run.font.size = Pt(8)

    dest.parent.mkdir(parents=True, exist_ok=True)
    if include_appendix:
        import tempfile
        from pathlib import Path as _Path
        from .appendix import append_to_docx
        with tempfile.TemporaryDirectory(prefix="pingtool_appendix_") as tmp:
            append_to_docx(doc, report, is_multi=True, tmp_dir=_Path(tmp))
            doc.save(str(dest))
    else:
        doc.save(str(dest))
