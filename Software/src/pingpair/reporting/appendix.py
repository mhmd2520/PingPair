"""Rich Analysis appendix for per-sweep reports.

Embedded in every saved docx / xlsx / pdf / txt (no longer optional;
the Save-Options-tab toggle was removed 2026-05-12). Each format gets the
richest visual the format can carry:

* **docx** — per-metric line charts (matplotlib PNG) + summary table +
  per-payload + per-bandwidth breakdown tables.
* **xlsx** — native openpyxl ``LineChart`` per metric + summary table +
  per-payload + per-bandwidth pivots in dedicated Analysis sheet.
* **pdf** — matplotlib PNG charts + tables (same as docx layout).
* **txt** — ASCII sparkline per metric + tables.

All renders are driven from a :class:`LoadedRun` via the
``from_run_report`` / ``from_multi_segment_report`` converters, so
the appendix uses the exact same aggregator the Analysis tab uses.
"""

from __future__ import annotations

import math
from pathlib import Path

from ..analysis import METRICS, fmt, run_stats
from ..analysis.chart_renderer import (
    ascii_sparkline,
    render_breakdown_chart,
    render_metric_chart,
)
from ..analysis.run_to_loaded import (
    from_multi_segment_report,
    from_run_report,
)


_APPENDIX_HEADERS: tuple[str, ...] = (
    "Metric", "Samples", "Min", "Avg", "Median", "Max", "Stdev",
)


# ---------------------------------------------------------------------------
# Shared aggregation helpers
# ---------------------------------------------------------------------------


def _stat_rows(loaded_run) -> list[list[str]]:
    """Per-metric summary stats table (one row per metric)."""
    rs = run_stats(loaded_run, case_filter=lambda _c: True)
    rows: list[list[str]] = []
    for metric in METRICS:
        ms = rs.by_metric[metric.code]
        d = metric.decimals
        rows.append([
            f"{metric.display} ({metric.unit})",
            str(ms.samples),
            fmt(ms.min, d),
            fmt(ms.avg, d),
            fmt(ms.median, d),
            fmt(ms.max, d),
            fmt(ms.stdev, d),
        ])
    return rows


def _flatten_cases(loaded_run) -> list:
    out = []
    for series in loaded_run.series:
        out.extend(series.cases)
    return out


def _breakdown_by_payload(loaded_run) -> list[list[str]]:
    """Avg-per-metric grouped by payload size. Header + body rows."""
    cases = _flatten_cases(loaded_run)
    buckets: dict[int, list] = {}
    for c in cases:
        buckets.setdefault(c.payload_bytes, []).append(c)
    rows: list[list[str]] = []
    for payload in sorted(buckets):
        vals = buckets[payload]
        row = [f"{payload} B", str(len(vals))]
        for metric in METRICS:
            samples = [
                getattr(c, metric.attr) for c in vals
                if getattr(c, metric.attr) is not None
                and math.isfinite(getattr(c, metric.attr))
            ]
            if samples:
                row.append(fmt(sum(samples) / len(samples), metric.decimals))
            else:
                row.append("—")
        rows.append(row)
    return rows


def _breakdown_by_bandwidth(loaded_run) -> list[list[str]]:
    cases = _flatten_cases(loaded_run)
    buckets: dict[int, list] = {}
    for c in cases:
        buckets.setdefault(c.bandwidth_mbps_pushed, []).append(c)
    rows: list[list[str]] = []
    for bw in sorted(buckets):
        vals = buckets[bw]
        row = [f"{bw} Mbps", str(len(vals))]
        for metric in METRICS:
            samples = [
                getattr(c, metric.attr) for c in vals
                if getattr(c, metric.attr) is not None
                and math.isfinite(getattr(c, metric.attr))
            ]
            if samples:
                row.append(fmt(sum(samples) / len(samples), metric.decimals))
            else:
                row.append("—")
        rows.append(row)
    return rows


def _breakdown_headers() -> list[str]:
    h = ["Bucket", "Cases"]
    for m in METRICS:
        h.append(f"Avg {m.display.split()[0]} ({m.unit})")
    return h


# ---------------------------------------------------------------------------
# DOCX
# ---------------------------------------------------------------------------


def append_to_docx(doc, report, *, is_multi: bool, tmp_dir: Path) -> None:
    """Append rich Analysis appendix to an open python-docx Document.

    ``tmp_dir`` must be a directory the caller created and will clean up
    *after* ``doc.save()`` returns. python-docx defers reading the
    embedded image files until save, so we cannot delete them inside
    this function.
    """
    from docx.shared import Inches

    loaded = (from_multi_segment_report(report)
              if is_multi else from_run_report(report))

    doc.add_page_break()
    doc.add_heading("Analysis appendix", level=1)
    doc.add_paragraph(
        "Per-metric summary statistics and per-case charts for this run, "
        "plus per-payload and per-bandwidth breakdowns. Open the Analysis "
        "tab in PingPair to overlay this run against past sweeps."
    )

    # Summary stats table
    doc.add_heading("Summary statistics", level=2)
    _docx_table(doc, list(_APPENDIX_HEADERS), _stat_rows(loaded))

    # Per-case charts
    doc.add_heading("Per-case charts", level=2)
    try:
        for metric in METRICS:
            png = tmp_dir / f"{metric.code}.png"
            result = render_metric_chart(loaded, metric, png)
            if result is None:
                doc.add_paragraph(
                    f"[{metric.display}: no plottable samples]"
                )
                continue
            doc.add_paragraph(metric.display, style="Heading 3")
            try:
                doc.add_picture(str(result), width=Inches(6.0))
            except Exception as exc:  # noqa: BLE001
                doc.add_paragraph(
                    f"[chart {metric.display} could not be embedded: {exc}]"
                )

        # Breakdown by payload
        doc.add_heading("Mean metrics by payload size", level=2)
        _docx_table(doc, _breakdown_headers(), _breakdown_by_payload(loaded))
        bp = tmp_dir / "by_payload.png"
        if render_breakdown_chart(
            loaded, METRICS[0], bp, by="payload"
        ) is not None:
            doc.add_paragraph("Avg throughput by payload", style="Heading 3")
            doc.add_picture(str(bp), width=Inches(6.0))

        # Breakdown by bandwidth
        doc.add_heading("Mean metrics by bandwidth", level=2)
        _docx_table(doc, _breakdown_headers(), _breakdown_by_bandwidth(loaded))
        bb = tmp_dir / "by_bandwidth.png"
        if render_breakdown_chart(
            loaded, METRICS[0], bb, by="bandwidth"
        ) is not None:
            doc.add_paragraph("Avg throughput by bandwidth", style="Heading 3")
            doc.add_picture(str(bb), width=Inches(6.0))
    finally:
        # The caller owns ``tmp_dir`` and is responsible for cleanup
        # after ``doc.save()`` returns — python-docx only reads the
        # embedded image paths at save time, so we can't delete them
        # inside this function.
        pass


def _docx_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    from docx.enum.table import WD_ALIGN_VERTICAL
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light List Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for r in hdr[i].paragraphs[0].runs:
            r.bold = True
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, cell in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = cell


# ---------------------------------------------------------------------------
# XLSX
# ---------------------------------------------------------------------------


def append_to_xlsx(wb, report, *, is_multi: bool) -> None:
    """Add a richer Analysis sheet (table + breakdowns + four native charts)."""
    from openpyxl.chart import LineChart, Reference
    from openpyxl.chart.marker import Marker
    from openpyxl.styles import Font

    loaded = (from_multi_segment_report(report)
              if is_multi else from_run_report(report))

    ws = wb.create_sheet("Analysis")
    ws["A1"] = "Analysis appendix"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A2"] = (
        "Per-metric stats, per-case data, and per-payload / per-bandwidth "
        "breakdowns. Charts on the right are native Excel charts you can "
        "re-pivot freely."
    )

    # Summary stats table (rows 4..)
    headers = list(_APPENDIX_HEADERS)
    for c, h in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=c, value=h)
        _style_header_cell(cell)
    rows = _stat_rows(loaded)
    for r_idx, row_data in enumerate(rows, start=5):
        for c_idx, cell in enumerate(row_data, start=1):
            ws.cell(row=r_idx, column=c_idx, value=cell)

    # Per-case data dump (chart source) — rows starting after the summary
    data_top_row = 4 + len(rows) + 2
    ws.cell(row=data_top_row, column=1, value="Case #").font = Font(bold=True)
    metric_cols: dict[str, int] = {}
    for i, metric in enumerate(METRICS):
        col = 2 + i
        ws.cell(
            row=data_top_row, column=col,
            value=f"{metric.display} ({metric.unit})",
        ).font = Font(bold=True)
        metric_cols[metric.code] = col

    flat_cases = _flatten_cases(loaded)
    for r_idx, case in enumerate(flat_cases, start=data_top_row + 1):
        ws.cell(row=r_idx, column=1, value=case.case_idx)
        for metric in METRICS:
            ws.cell(
                row=r_idx, column=metric_cols[metric.code],
                value=getattr(case, metric.attr),
            )

    # Per-payload breakdown table (further down, left column). Compute
    # each breakdown once — it flattens + buckets every case, so the old
    # in-line recomputation inside the row-offset math was wasted work.
    payload_rows = _breakdown_by_payload(loaded)
    bandwidth_rows = _breakdown_by_bandwidth(loaded)
    bp_top = data_top_row + len(flat_cases) + 3
    ws.cell(row=bp_top, column=1, value="Mean by payload size").font = Font(bold=True)
    bp_headers = _breakdown_headers()
    for c, h in enumerate(bp_headers, start=1):
        _style_header_cell(ws.cell(row=bp_top + 1, column=c, value=h))
    for r_idx, row_data in enumerate(payload_rows, start=bp_top + 2):
        for c_idx, cell in enumerate(row_data, start=1):
            ws.cell(row=r_idx, column=c_idx, value=cell)

    bb_top = bp_top + 2 + len(payload_rows) + 2
    ws.cell(row=bb_top, column=1, value="Mean by bandwidth").font = Font(bold=True)
    for c, h in enumerate(bp_headers, start=1):
        _style_header_cell(ws.cell(row=bb_top + 1, column=c, value=h))
    for r_idx, row_data in enumerate(bandwidth_rows, start=bb_top + 2):
        for c_idx, cell in enumerate(row_data, start=1):
            ws.cell(row=r_idx, column=c_idx, value=cell)

    if not flat_cases:
        # No case data at all (sweep aborted before any case ran). Drop a
        # clear placeholder so the user knows the missing charts aren't
        # a rendering bug. (Task K, 2026-05-12.)
        from openpyxl.styles import Font as _Font
        placeholder_cell = ws.cell(
            row=3, column=9,
            value="No chart data — this sweep recorded 0 cases.",
        )
        placeholder_cell.font = _Font(italic=True, color="FF888888")
        return

    # Native line charts anchored to the right of the stats table.
    # openpyxl LineChart defaults hide numeric tick labels on both axes
    # unless we explicitly configure them — the symptom is "title shows
    # but the X/Y numbers are blank in Excel". Forcing
    # delete=False / majorTickMark="out" / tickLblPos="nextTo" fixes it.
    # (Task H, 2026-05-12.) Per-metric data check (Task K) skips a
    # chart whose metric has no non-None values, so we don't render an
    # axes-only empty box.
    data_bottom_row = data_top_row + len(flat_cases)
    placeholder_anchor = 3
    for chart_idx, metric in enumerate(METRICS):
        # Skip if this metric column has no plottable values.
        has_data = any(
            getattr(c, metric.attr) is not None for c in flat_cases
        )
        if not has_data:
            from openpyxl.styles import Font as _Font
            note = ws.cell(
                row=placeholder_anchor + chart_idx * 18, column=9,
                value=(
                    f"{metric.display}: no data (every case missing this metric)."
                ),
            )
            note.font = _Font(italic=True, color="FF888888")
            continue
        chart = LineChart()
        chart.title = f"{metric.display} ({metric.unit})"
        chart.y_axis.title = metric.unit
        chart.x_axis.title = "Case #"
        chart.height = 8
        chart.width = 16
        chart.legend = None
        chart.style = 2  # solid line + visible markers

        # Force both axes to render with numeric tick labels.
        chart.x_axis.delete = False
        chart.y_axis.delete = False
        chart.x_axis.majorTickMark = "out"
        chart.y_axis.majorTickMark = "out"
        chart.x_axis.tickLblPos = "nextTo"
        chart.y_axis.tickLblPos = "nextTo"
        # crossAx must reference the OTHER axis; openpyxl sets defaults
        # but they sometimes confuse Excel into hiding the tick labels.
        chart.x_axis.crossAx = chart.y_axis.axId
        chart.y_axis.crossAx = chart.x_axis.axId

        ref = Reference(
            ws,
            min_col=metric_cols[metric.code],
            min_row=data_top_row,
            max_col=metric_cols[metric.code],
            max_row=data_bottom_row,
        )
        chart.add_data(ref, titles_from_data=True)
        cat_ref = Reference(
            ws, min_col=1, min_row=data_top_row + 1,
            max_row=data_bottom_row,
        )
        chart.set_categories(cat_ref)
        # Explicit per-series markers. Without them a SINGLE-case run (one data
        # point) draws a zero-length line with nothing visible — the chart looks
        # empty even though the cell has data. (chart.style=2 does NOT reliably
        # emit markers in openpyxl's output.) Markers also make multi-point
        # series easier to read. (#2, 2026-06-01.)
        for series in chart.series:
            series.marker = Marker(symbol="circle", size=6)
            series.smooth = False
        anchor_row = 3 + chart_idx * 18
        ws.add_chart(chart, f"I{anchor_row}")


def _style_header_cell(cell) -> None:
    from openpyxl.styles import Font, PatternFill, Alignment
    cell.font = Font(bold=True, color="FFFFFFFF")
    cell.fill = PatternFill(
        start_color="FF2A3F66",
        end_color="FF2A3F66",
        fill_type="solid",
    )
    cell.alignment = Alignment(horizontal="center")


# ---------------------------------------------------------------------------
# PDF
# ---------------------------------------------------------------------------


def append_to_pdf(flowables, report, *, is_multi: bool, tmp_dir: Path) -> None:
    """Append rich Analysis appendix to a reportlab flowables list.

    ``tmp_dir`` must be a directory the caller created and will clean up
    *after* ``doc.build()`` returns. reportlab's Image flowable defers
    reading PNGs until build time.
    """
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import (
        Image,
        PageBreak,
        Paragraph,
        Spacer,
    )

    loaded = (from_multi_segment_report(report)
              if is_multi else from_run_report(report))
    styles = getSampleStyleSheet()
    flowables.append(PageBreak())
    flowables.append(Paragraph("<b>Analysis appendix</b>", styles["Heading1"]))
    flowables.append(Paragraph(
        "Per-metric stats and per-case charts for this run, plus "
        "per-payload and per-bandwidth breakdowns.",
        styles["Normal"],
    ))
    flowables.append(Spacer(1, 6))

    headers = list(_APPENDIX_HEADERS)
    rows = _stat_rows(loaded)
    flowables.append(_pdf_table(headers, rows))
    flowables.append(Spacer(1, 10))

    for metric in METRICS:
        png = tmp_dir / f"{metric.code}.png"
        result = render_metric_chart(loaded, metric, png)
        if result is None:
            flowables.append(Paragraph(
                f"[{metric.display}: no plottable samples]",
                styles["Normal"],
            ))
            continue
        flowables.append(Paragraph(
            f"<b>{metric.display}</b>", styles["Heading3"]
        ))
        flowables.append(Image(str(result), width=22 * cm, height=10 * cm))
        flowables.append(Spacer(1, 6))

    flowables.append(PageBreak())
    flowables.append(Paragraph(
        "<b>Mean metrics by payload size</b>", styles["Heading2"]))
    flowables.append(_pdf_table(_breakdown_headers(),
                                _breakdown_by_payload(loaded)))
    bp = tmp_dir / "by_payload.png"
    if render_breakdown_chart(loaded, METRICS[0], bp, by="payload") is not None:
        flowables.append(Spacer(1, 6))
        flowables.append(Image(str(bp), width=22 * cm, height=10 * cm))

    flowables.append(Spacer(1, 10))
    flowables.append(Paragraph(
        "<b>Mean metrics by bandwidth</b>", styles["Heading2"]))
    flowables.append(_pdf_table(_breakdown_headers(),
                                _breakdown_by_bandwidth(loaded)))
    bb = tmp_dir / "by_bandwidth.png"
    if render_breakdown_chart(loaded, METRICS[0], bb, by="bandwidth") is not None:
        flowables.append(Spacer(1, 6))
        flowables.append(Image(str(bb), width=22 * cm, height=10 * cm))


def _pdf_table(headers, rows):
    from reportlab.lib import colors
    from reportlab.platypus import Table, TableStyle
    table = Table([headers] + rows, repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2a3f66")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, 0), 8),
        ("FONTSIZE", (0, 1), (-1, -1), 8),
        ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("GRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#cccccc")),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.white, colors.HexColor("#f7f7f9")]),
    ]))
    return table


# ---------------------------------------------------------------------------
# TXT
# ---------------------------------------------------------------------------


def append_to_txt(lines: list[str], report, *, is_multi: bool) -> None:
    loaded = (from_multi_segment_report(report)
              if is_multi else from_run_report(report))
    lines.append("")
    lines.append("Analysis appendix")
    lines.append("=" * len("Analysis appendix"))
    lines.append("")
    lines.append(
        "Per-metric summary statistics plus per-case ASCII sparklines "
        "and per-payload / per-bandwidth breakdowns. For the rich "
        "visual version, see the docx / xlsx / pdf alongside this file."
    )
    lines.append("")

    # Summary stats
    headers = list(_APPENDIX_HEADERS)
    rows = _stat_rows(loaded)
    lines.extend(_txt_table(headers, rows))
    lines.append("")

    # Sparklines per metric
    lines.append("Per-case sparklines (one block per case, "
                 "min↦blank, max↦█)")
    lines.append("-" * 64)
    flat = _flatten_cases(loaded)
    for metric in METRICS:
        values = [getattr(c, metric.attr) for c in flat]
        line = ascii_sparkline(values)
        lo = min(
            (v for v in values if v is not None and math.isfinite(v)),
            default=None,
        )
        hi = max(
            (v for v in values if v is not None and math.isfinite(v)),
            default=None,
        )
        rng = (f"min={fmt(lo, metric.decimals)} "
               f"max={fmt(hi, metric.decimals)} {metric.unit}"
               if lo is not None else "(no samples)")
        lines.append(
            f"  {metric.display:<14} | {line}  [{rng}]"
        )
    lines.append("")

    # Breakdown tables
    lines.append("Mean metrics by payload size")
    lines.append("-" * len("Mean metrics by payload size"))
    lines.extend(_txt_table(_breakdown_headers(),
                            _breakdown_by_payload(loaded)))
    lines.append("")
    lines.append("Mean metrics by bandwidth")
    lines.append("-" * len("Mean metrics by bandwidth"))
    lines.extend(_txt_table(_breakdown_headers(),
                            _breakdown_by_bandwidth(loaded)))


def _txt_table(headers: list[str], rows: list[list[str]]) -> list[str]:
    if not headers:
        return []
    widths = [len(h) for h in headers]
    for row in rows:
        for col, cell in enumerate(row):
            if col < len(widths):
                widths[col] = max(widths[col], len(cell))

    def _row(values: list[str]) -> str:
        return "  ".join(v.ljust(widths[i]) for i, v in enumerate(values))

    out = [_row(headers), _row(["-" * w for w in widths])]
    for row in rows:
        out.append(_row(row))
    return out
