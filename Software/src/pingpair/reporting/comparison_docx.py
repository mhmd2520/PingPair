"""DOCX writer for the Analysis-tab comparison report (#12).

Layout:

1. Title heading + generated-at line.
2. "Runs included" table (one row per run).
3. "Filter applied" bullet list (skipped when the filter is at defaults).
4. "Summary statistics" table — one row per run, columns by metric.
5. Embedded charts (one figure per metric) when PNGs were attached.
6. "Per-case delta" table — only when exactly 2 runs are included.
7. Optional notes section.

Uses python-docx (already a PingPair dep for the per-sweep writer).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches, Pt

from ..analysis import METRICS, fmt, fmt_delta
from ..analysis.comparison import ComparisonReport
from ._logo import add_docx_logo
from .run_report import fmt_duration


def _add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    """Add a styled table with one header row and ``rows`` body rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light List Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for r_idx, row_data in enumerate(rows, start=1):
        for c_idx, cell_text in enumerate(row_data):
            table.rows[r_idx].cells[c_idx].text = cell_text


def write_comparison_docx(report: ComparisonReport, dest: Path) -> None:
    """Write ``report`` to ``dest`` as a .docx file."""
    doc = Document()

    # Title
    add_docx_logo(doc)
    doc.add_heading(report.title, level=0)
    p = doc.add_paragraph()
    p.add_run(
        f"Generated: {report.generated_at.strftime('%Y-%m-%d %H:%M:%S')}"
    ).font.size = Pt(10)
    p.add_run(f"   ·   Runs compared: {report.run_count}").font.size = Pt(10)

    if report.notes:
        doc.add_heading("Notes", level=1)
        for line in report.notes.splitlines():
            doc.add_paragraph(line)

    # Runs included
    doc.add_heading("Runs included", level=1)
    run_rows: list[list[str]] = []
    for run in report.runs:
        when = (
            run.started_at.strftime("%Y-%m-%d %H:%M")
            if run.started_at is not None
            else "?"
        )
        run_rows.append([
            run.display_label,
            when,
            fmt_duration(run.duration_s),
            "multi" if run.is_multi_segment else "single",
            f"{run.cases_ok}/{run.cases_total} ok",
        ])
    _add_table(
        doc,
        ["Run", "Started", "Duration", "Type", "Cases"],
        run_rows,
    )

    # Filter applied
    if not report.filter_description.is_default:
        doc.add_heading("Filter applied", level=2)
        for line in report.filter_description.lines():
            doc.add_paragraph(line, style="List Bullet")

    # Summary stats
    doc.add_heading("Summary statistics", level=1)
    headers = ["Run"]
    for metric in METRICS:
        short = metric.display.split()[0]
        headers.extend([
            f"{short} min", f"{short} avg", f"{short} max",
        ])
    stat_rows: list[list[str]] = []
    for run, rs in zip(report.runs, report.per_run_stats, strict=False):
        row = [run.display_label]
        for metric in METRICS:
            ms = rs.by_metric[metric.code]
            row.extend([fmt(ms.min), fmt(ms.avg), fmt(ms.max)])
        stat_rows.append(row)
    _add_table(doc, headers, stat_rows)

    # Charts
    if report.chart_pngs:
        doc.add_heading("Charts", level=1)
        for metric in METRICS:
            png = report.chart_pngs.get(metric.code)
            if png is None or not png.exists():
                continue
            doc.add_paragraph(metric.display, style="Heading 3")
            try:
                doc.add_picture(str(png), width=Inches(6.0))
            except Exception as exc:  # noqa: BLE001
                doc.add_paragraph(
                    f"[chart {metric.display} could not be embedded: {exc}]"
                )

    # Diff
    if report.has_diff_section:
        a = report.runs[1]
        b = report.runs[0]
        doc.add_heading("Per-case delta (B − A)", level=1)
        doc.add_paragraph(
            f"A (older): {a.display_label}    B (newer): {b.display_label}"
        )
        diff_headers = ["Case", "Payload (B)", "BW (Mbps)"]
        for metric in METRICS:
            short = metric.display.split()[0]
            diff_headers.extend([
                f"A {short}", f"B {short}", f"Δ {short}",
            ])
        diff_rows: list[list[str]] = []
        for row in report.per_case_diff_rows:
            r = [
                str(row.case_idx),
                str(row.payload_bytes),
                str(row.bandwidth_mbps_pushed),
            ]
            for metric in METRICS:
                r.extend([
                    fmt(row.a_value[metric.code]),
                    fmt(row.b_value[metric.code]),
                    fmt_delta(row.delta[metric.code]),
                ])
            diff_rows.append(r)
        _add_table(doc, diff_headers, diff_rows)

    doc.save(str(dest))
