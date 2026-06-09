"""Shared report-header logo helper.

Embeds the canonical PingPair mark (:data:`pingpair.branding.ICON_SVG`) at the
top of the Word / PDF report headers. The raster comes from
:func:`pingpair.branding.logo_png_bytes` — the QImage paint path, which needs no
running QApplication, so it works from the report worker thread and in headless
unit tests.

Every entry point is failure-tolerant: if Qt is unavailable or rendering fails
for any reason, the logo is simply omitted so report generation never breaks.
``branding`` (and therefore PySide6) is imported lazily inside the helpers, which
keeps ``import pingpair.reporting`` free of a top-level Qt import.
"""

from __future__ import annotations

import contextlib
import io

# Source raster resolution for the embedded mark. The logo is *displayed* small
# (≈18 mm / 0.7 in), but we rasterize the vector at a much larger pixel size so
# it stays crisp under heavy PDF zoom and on print (768 px ≈ 1080 DPI at that
# physical size, yet only ~20 KB for this simple mark). Display width is set per
# format below — this only controls pixel density, not how big it looks.
_LOGO_RENDER_PX = 768


def _logo_png(size: int = _LOGO_RENDER_PX) -> bytes | None:
    """Return PNG bytes of the mark, or ``None`` if rendering isn't possible."""
    try:
        from ..branding import logo_png_bytes

        return logo_png_bytes(size)
    except Exception:
        # A missing/broken renderer (no Qt, headless edge case) must never break
        # a report — fall back to "no logo".
        return None


def add_docx_logo(doc, *, width_in: float = 0.7) -> None:
    """Add the centred PingPair mark above the title (no-op if rendering fails)."""
    png = _logo_png()
    if png is None:
        return

    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        para.add_run().add_picture(io.BytesIO(png), width=Inches(width_in))
    except Exception:
        # Near-impossible (the bytes are already valid), but roll back the empty
        # paragraph so a failed embed can't leave a stray blank line.
        with contextlib.suppress(Exception):
            para._element.getparent().remove(para._element)


def pdf_logo_flowable(*, size_mm: float = 18.0):
    """Return a centred reportlab ``Image`` flowable, or ``None`` on failure."""
    png = _logo_png()
    if png is None:
        return None

    from reportlab.lib.units import mm
    from reportlab.platypus import Image

    try:
        # reportlab eagerly reads file-like inputs at construction (it forces
        # lazy=0 internally for a BytesIO), so the buffer needn't survive until
        # build time; passing lazy=0 just makes that explicit.
        img = Image(io.BytesIO(png), width=size_mm * mm, height=size_mm * mm, lazy=0)
    except Exception:
        return None
    img.hAlign = "CENTER"
    return img


# Standard gap below the mark in a PDF report header — one source so every PDF
# writer's logo spacing stays identical (~4 mm).
_PDF_LOGO_GAP_MM = 4.0


def pdf_logo_block() -> list[object]:
    """Return ``[logo, spacer]`` to drop at the top of a PDF report's flowables,
    or ``[]`` if the logo can't be rendered. Centralises the mark plus its standard
    header gap so every PDF writer (run report, multi-segment, comparison) inserts
    an identical block."""
    logo = pdf_logo_flowable()
    if logo is None:
        return []

    from reportlab.lib.units import mm
    from reportlab.platypus import Spacer

    return [logo, Spacer(1, _PDF_LOGO_GAP_MM * mm)]
